from flask import Flask, request, render_template, send_file, redirect, url_for, flash
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from dotenv import load_dotenv
from datetime import datetime
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import os

# Load environment variables from .env file
load_dotenv()
DATABASE_URL = os.getenv("DATABASE_URL")

# Ensure the database URL is available
if not DATABASE_URL:
    raise ValueError("DATABASE_URL environment variable is not set")

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "your_secret_key")  # You need a secret key for flash messages

# Database setup with SSL if necessary
engine = create_engine(DATABASE_URL, connect_args={"sslmode": "require"})
Base = declarative_base()
Session = sessionmaker(bind=engine)

# Define the Inquiry model
class Inquiry(Base):
    __tablename__ = 'inquiries'
    id = Column(Integer, primary_key=True)
    client_name = Column(String(100))
    phone_number = Column(String(20))
    client_inquiry = Column(Text)
    booking_date = Column(DateTime)
    resolution = Column(Text)
    date_called = Column(DateTime)
    feedback = Column(Text)
    conversion = Column(String(20))
    created_at = Column(DateTime, default=datetime.utcnow)

# Route to create database tables (temporary, remove after setup)
@app.route('/create-db')
def create_db():
    Base.metadata.create_all(bind=engine)
    return "Database tables created!"

# Route to display the form
@app.route('/')
def form():
    return render_template("index.html")

# Route to handle form submission and save inquiry
@app.route('/save-inquiry', methods=['POST'])
def save_inquiry():
    # Get form data
    client_name = request.form['client-name']
    phone_number = request.form['phone-number']
    client_inquiry = request.form['client-inquiry']
    booking_date = datetime.strptime(request.form['booking-date'], '%Y-%m-%d')  # Convert to datetime
    resolution = request.form['resolution']
    date_called = datetime.strptime(request.form['date-called'], '%Y-%m-%d') if request.form['date-called'] else None  # Convert to datetime
    feedback = request.form['client-feedback']
    conversion = request.form['conversion']
    
    # Use context manager for session handling
    with Session() as session:
        new_inquiry = Inquiry(
            client_name=client_name,
            phone_number=phone_number,
            client_inquiry=client_inquiry,
            booking_date=booking_date,
            resolution=resolution,
            date_called=date_called,
            feedback=feedback,
            conversion=conversion
        )
        session.add(new_inquiry)
        session.commit()

    # Flash a success message and redirect to the inquiries page
    flash('Inquiry submitted successfully!', 'success')
    return redirect(url_for('inquiries'))

@app.route('/inquiries')
def inquiries():
    with Session() as session:
        inquiries = session.query(Inquiry).order_by(Inquiry.created_at.desc()).all()
        
        # Render inquiries in an HTML page and link the styles.css
        return render_template('inquiries.html', inquiries=inquiries)


# Route to generate PDF for a specific inquiry
@app.route('/download-pdf/<int:inquiry_id>')
def download_pdf(inquiry_id):
    with Session() as session:
        inquiry = session.query(Inquiry).filter(Inquiry.id == inquiry_id).first()
        
        # Check if the inquiry exists
        if not inquiry:
            return "Inquiry not found", 404
        
        # Create a PDF in memory
        pdf_buffer = BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        c.setFont("Helvetica", 12)
        y_position = 750
        
        # Add title
        c.drawString(200, y_position, f"Inquiry #{inquiry.id} - {inquiry.client_name}")
        y_position -= 20
        
        # Add inquiry details to the PDF
        c.drawString(30, y_position, f"Client Name: {inquiry.client_name}")
        y_position -= 15
        c.drawString(30, y_position, f"Phone Number: {inquiry.phone_number}")
        y_position -= 15
        c.drawString(30, y_position, f"Inquiry: {inquiry.client_inquiry}")
        y_position -= 15
        c.drawString(30, y_position, f"Booking Date: {inquiry.booking_date}")
        y_position -= 15
        c.drawString(30, y_position, f"Resolution: {inquiry.resolution}")
        
        # Finalize the PDF
        c.save()

        # Return the PDF file
        pdf_buffer.seek(0)
        return send_file(pdf_buffer, as_attachment=True, download_name=f"inquiry_{inquiry.id}.pdf", mimetype="application/pdf")

# Route to generate Word file for a specific inquiry
@app.route('/download-docx/<int:inquiry_id>')
def download_docx(inquiry_id):
    with Session() as session:
        inquiry = session.query(Inquiry).filter(Inquiry.id == inquiry_id).first()
        
        # Check if the inquiry exists
        if not inquiry:
            return "Inquiry not found", 404
        
        # Create a Word document in memory
        doc = Document()
        doc.add_heading(f"Inquiry #{inquiry.id} - {inquiry.client_name}", 0)
        
        # Add inquiry details to the Word document
        doc.add_paragraph(f"Client Name: {inquiry.client_name}")
        doc.add_paragraph(f"Phone Number: {inquiry.phone_number}")
        doc.add_paragraph(f"Inquiry: {inquiry.client_inquiry}")
        doc.add_paragraph(f"Booking Date: {inquiry.booking_date}")
        doc.add_paragraph(f"Resolution: {inquiry.resolution}")
        
        # Save the Word document to a BytesIO buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        # Return the Word file
        return send_file(doc_buffer, as_attachment=True, download_name=f"inquiry_{inquiry.id}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Run the app
if __name__ == '__main__':
    port = int(os.getenv("PORT", 8080))  # Use PORT env variable if available, otherwise default to 8080
    app.run(host='0.0.0.0', port=port, debug=True)
